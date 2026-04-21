from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import httpx
import os
import shutil
import uuid
from docx import Document

app = FastAPI(title="DocTranslate API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En production, mets ton domaine frontend ici
    allow_methods=["*"],
    allow_headers=["*"],
)

DEEPL_API_KEY = os.getenv("DEEPL_API_KEY")  # Mets ta clé dans les variables d'environnement
DEEPL_URL = "https://api-free.deepl.com/v2/translate"  # URL Free tier


async def translate_text(text: str, target_lang: str, source_lang: str = "FR") -> str:
    if not text.strip():
        return text
    # Adapter le code langue pour MyMemory
    lang_map = {"EN-GB": "EN", "PT-BR": "PT"}
    src = lang_map.get(source_lang, source_lang)
    tgt = lang_map.get(target_lang, target_lang)
    url = "https://api.mymemory.translated.net/get"
    params = {"q": text, "langpair": f"{src}|{tgt}"}
    async with httpx.AsyncClient() as client:
        response = await client.get(url, params=params)
        data = response.json()
        return data["responseData"]["translatedText"]

@app.post("/translate")
async def translate_document(
    file: UploadFile = File(...),
    target_lang: str = Form(...),
    source_lang: str = Form(default=None),
):
    """
    Reçoit un fichier .docx, traduit tout le texte avec DeepL,
    retourne un nouveau .docx avec la mise en page conservée.
    """
    # Vérifier que c'est bien un .docx
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Seuls les fichiers .docx sont acceptés.")

    # Sauvegarder le fichier uploadé temporairement
    temp_id = str(uuid.uuid4())
    import tempfile
    temp_dir = tempfile.gettempdir()
    input_path = os.path.join(temp_dir, f"{temp_id}_input.docx")
    output_path = os.path.join(temp_dir, f"{temp_id}_output.docx")

    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Ouvrir le document
    doc = Document(input_path)

    # --- Traduire tous les paragraphes ---
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Construire le texte complet du paragraphe
            full_text = paragraph.text
            translated = await translate_text(full_text, target_lang, source_lang)

            # Réécrire le texte dans le premier run, vider les autres
            # (pour conserver le style du premier run)
            if paragraph.runs:
                paragraph.runs[0].text = translated
                for run in paragraph.runs[1:]:
                    run.text = ""

    # --- Traduire les tableaux ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        translated = await translate_text(paragraph.text, target_lang, source_lang)
                        if paragraph.runs:
                            paragraph.runs[0].text = translated
                            for run in paragraph.runs[1:]:
                                run.text = ""

    # Sauvegarder le document traduit
    doc.save(output_path)

    # Nettoyer le fichier d'entrée
    os.remove(input_path)

    # Retourner le fichier traduit
    translated_filename = f"traduit_{file.filename}"
    return FileResponse(
        path=output_path,
        filename=translated_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/health")
def health_check():
    return {"status": "ok", "message": "DocTranslate API is running"}